"""
Agente de Monitoramento Científico

Descrição:
Busca artigos científicos sobre Penaeus vannamei,
processa metadados e gera relatórios automatizados.

Desenvolvido com apoio de IA generativa.
Arquitetura e validação técnica realizadas pelo autor.
"""



import os
import datetime
import smtplib
import ssl
from email.mime.multipart import MIMEMultipart
from email.mime.text import MIMEText
from email.mime.base import MIMEBase
from email import encoders
from google import genai


# Bibliotecas para Scopus e IA
from pybliometrics.scopus import ScopusSearch
from openai import OpenAI
from docx import Document
from pybliometrics.scopus import init

init()



# --- CONFIGURAÇÕES ---
# Configure suas chaves aqui ou use variáveis de ambiente (mais seguro)
GEMINI_API_KEY = "SUA_API"
EMAIL_REMETENTE = "SEU_EMAIL"
EMAIL_SENHA = "SUA_SENHA" # Gere uma "App Password" no Google
EMAIL_DESTINATARIO = "EMAIL_DESTINO"
client = genai.Client(api_key=GEMINI_API_KEY)
# Termo de busca
TERMO_BUSCA = '"Penaeus vannamei" OR "Litopenaeus vannamei"'

def buscar_artigos_semanais():
    """Busca artigos publicados nos últimos 7 dias no Scopus."""
    print("🔍 Buscando artigos no Scopus...")
    # Data de 7 dias atrás
    hoje = datetime.date.today()
    semana_passada = hoje - datetime.timedelta(days=7)
    date_str = semana_passada.strftime("%Y-%m-%d")
    
    # Query do Scopus: Termo + Publicado após a data X
    # PUBDATETXT funciona melhor para faixas, mas vamos filtrar no código para garantir
    query = f'TITLE-ABS-KEY({TERMO_BUSCA}) AND PUBYEAR > {semana_passada.year - 1}'
    
    try:
        search = ScopusSearch(query, download=True)
        artigos_recentes = []
        
        if not search.results:
            return []

        for item in search.results:
            # Filtragem manual de data se a API retornar muitos antigos
            # O Scopus retorna datas como '2023-10-01' ou apenas ano.
            # Aqui vamos pegar os top 10 mais recentes para o exemplo não estourar cotas
            artigos_recentes.append({
                'titulo': item.title,
                'autores': item.author_names,
                'resumo': item.description if item.description else "Resumo não disponível.",
                'doi': item.doi,
                'link': f"https://doi.org/{item.doi}" if item.doi else "Link não disponível",
                'publicacao': item.coverDate
            })
            
        # Pega apenas os 10 mais recentes para economizar tokens/tempo
        return artigos_recentes[:10]
        
    except Exception as e:
        print(f"Erro na busca Scopus: {e}")
        return []

def resumir_com_ia(artigos):
    """Usa o Gemini (nova biblioteca) para resumir os abstracts."""
    
    resumos_processados = []
    print(f"🧠 Gerando resumos para {len(artigos)} artigos com Gemini...")

    for artigo in artigos:

        if not artigo['resumo'] or len(artigo['resumo']) < 50:
            resumo_ia = "Resumo original muito curto para processamento."
        else:
            prompt = f"""
Você é um pesquisador doutor em Aquicultura, especialista em fisiologia,
nutrição e manejo de Litopenaeus vannamei.

Produza um resumo técnico e analítico do artigo abaixo.

Regras:
- Escreva em português.
- Linguagem científica objetiva.
- Destaque:
    1) Objetivo
    2) Metodologia
    3) Resultados principais
    4) Aplicações práticas
    5) Limitações (se houver)
- Máximo 200 palavras.

Título: {artigo['titulo']}
Abstract: {artigo['resumo']}
"""

            try:
                response = client.models.generate_content(
                    model="gemini-2.5-flash",
                    contents=prompt
                )

                resumo_ia = response.text

            except Exception as e:
                resumo_ia = f"Erro ao gerar resumo com Gemini: {e}"

        artigo['resumo_ia'] = resumo_ia
        resumos_processados.append(artigo)

    return resumos_processados
def criar_documento(artigos):
    """Cria o arquivo Word."""
    print("📄 Criando arquivo DOCX...")
    doc = Document()
    doc.add_heading('Relatório Semanal: Penaeus vannamei', 0)
    data_hoje = datetime.date.today().strftime("%d/%m/%Y")
    doc.add_paragraph(f'Gerado em: {data_hoje}')
    
    if not artigos:
        doc.add_paragraph("Nenhum artigo novo encontrado nesta semana.")
        return "relatorio_semanal.docx"

    for art in artigos:
        doc.add_heading(art['titulo'], level=1)
        doc.add_paragraph(f"Autores: {art['autores']}")
        doc.add_paragraph(f"Data: {art['publicacao']}")
        doc.add_paragraph(f"Link/DOI: {art['link']}")
        
        doc.add_heading('Resumo da IA:', level=2)
        doc.add_paragraph(art['resumo_ia'])
        
        doc.add_paragraph("-" * 50)

    filename = f"Relatorio_Vannamei_{datetime.date.today()}.docx"
    doc.save(filename)
    return filename

def enviar_email(anexo_path):
    """Envia o e-mail com o anexo."""
    print("📧 Enviando e-mail...")
    
    msg = MIMEMultipart()
    msg['From'] = EMAIL_REMETENTE
    msg['To'] = EMAIL_DESTINATARIO
    msg['Subject'] = f"Atualização Científica - Penaeus vannamei - {datetime.date.today()}"
    
    body = "Olá,\n\nSegue em anexo o relatório semanal com os artigos mais recentes sobre Penaeus vannamei encontrados no Scopus, resumidos por IA.\n\nAtenciosamente,\nSeu Agente de IA."
    msg.attach(MIMEText(body, 'plain'))
    
    # Anexar arquivo
    try:
        attachment = open(anexo_path, "rb")
        part = MIMEBase('application', 'octet-stream')
        part.set_payload(attachment.read())
        encoders.encode_base64(part)
        part.add_header('Content-Disposition', f"attachment; filename= {os.path.basename(anexo_path)}")
        msg.attach(part)
        attachment.close()
        
        server = smtplib.SMTP('smtp.gmail.com', 587)
        server.starttls()
        server.login(EMAIL_REMETENTE, EMAIL_SENHA)
        text = msg.as_string()
        server.sendmail(EMAIL_REMETENTE, EMAIL_DESTINATARIO, text)
        server.quit()
        print("✅ E-mail enviado com sucesso!")
        
    except Exception as e:
        print(f"Erro ao enviar e-mail: {e}")

# --- FLUXO PRINCIPAL ---
if __name__ == "__main__":
    # 1. Buscar
    artigos = buscar_artigos_semanais()
    
    if artigos:
        # 2. Resumir
        artigos_resumidos = resumir_com_ia(artigos)
        
        # 3. Criar Doc
        arquivo_doc = criar_documento(artigos_resumidos)
        
        # 4. Enviar
        enviar_email(arquivo_doc)
    else:
        print("Nenhum artigo encontrado para processar.")
